import os
from slack_sdk import WebClient
from slack_sdk.errors import SlackApiError
from dotenv import load_dotenv
from datetime import datetime
import threading
from elevenlabs import ElevenLabs
import logging
from docx import Document

class SlackUploader:
    def __init__(self):
        self.client = WebClient(token=os.getenv('SLACK_BOT_TOKEN'))
        self.channel = "C084DBCMQR4"
        
    def upload_story(self, content, topic, current_date):
        try:
            message_text = f"{topic['situation']}.\n{content}"
            
            blocks = [
                {
                    "type": "section",
                    "text": {
                        "type": "mrkdwn",
                        "text": message_text
                    }
                },
                {
                    "type": "section",
                    "text": {
                        "type": "mrkdwn",
                        "text": "`결과물 생성` 버튼을 클릭하거나 스레드에 `수정본:`으로 시작하는 수정본을 업로드하세요."
                    }
                },
                {
                    "type": "actions",
                    "elements": [
                        {
                            "type": "button",
                            "text": {
                                "type": "plain_text",
                                "text": "결과물 생성",
                                "emoji": True
                            },
                            "action_id": "generate_content",
                            "value": "generate"
                        }
                    ]
                }
            ]

            response = self.client.chat_postMessage(
                channel=self.channel,
                blocks=blocks,
                text="새로운 에피소드가 업로드되었습니다."
            )
            
            return response
            
        except SlackApiError as e:
            print(f"Slack API Error: {e.response['error']}")
            return None

    def generate_content(self, text, thread_ts):
        """텍스트를 docx와 오디오로 변환하여 업로드"""
        print(f"Starting generate_content with text length: {len(text)} and thread_ts: {thread_ts}")
        try:
            # docx 생성 및 업로드
            docx_file = self.create_docx("에피소드", text)
            print(f"Created docx file: {docx_file}")

            docx_response = self.client.files_upload_v2(
                channel=self.channel,
                thread_ts=thread_ts,
                file=docx_file,
                title="텍스트 버전",
                initial_comment="📄 텍스트 버전이 준비되었습니다."
            )
            print(f"Uploaded docx file: {docx_response}")
            
            try:
                os.remove(docx_file)
                print(f"Removed temp docx file: {docx_file}")
            except Exception as e:
                print(f"Error removing temp docx file: {e}")

            # 오디오 생성 및 업로드
            audio_file = self.generate_audio(text)
            print(f"Created audio file: {audio_file}")
            
            audio_response = self.client.files_upload_v2(
                channel=self.channel,
                thread_ts=thread_ts,
                file=audio_file,
                title="음성 버전",
                initial_comment="🎧 음성 버전이 준비되었습니다."
            )
            print(f"Uploaded audio file: {audio_response}")
            
            try:
                os.remove(audio_file)
                print(f"Removed temp audio file: {audio_file}")
            except Exception as e:
                print(f"Error removing temp audio file: {e}")

        except Exception as e:
            print(f"Error in generate_content: {e}")
            self.client.chat_postMessage(
                channel=self.channel,
                thread_ts=thread_ts,
                text=f"⚠️ 결과물 생성 중 오류가 발생했습니다: {str(e)}"
            )

    def create_docx(self, title, content):
        """docx 파일 생성"""
        print(f"Creating docx with title: {title}")
        try:
            doc = Document()
            doc.add_heading(title, 0)
            doc.add_paragraph(content)
            
            filename = f"episode_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)
            print(f"Successfully created docx: {filename}")
            return filename
            
        except Exception as e:
            print(f"Error creating docx: {e}")
            raise

    def generate_audio(self, text: str) -> str:
        """에세이 텍스트를 오디오로 변환"""
        try:
            client = ElevenLabs(api_key=os.getenv("ELEVENLABS_API_KEY"))
            
            # 음성 생성
            audio_stream = client.text_to_speech.convert(
                voice_id="UvSpBkqvdJNLIrGN8kdQ",
                text=text,
                model_id="eleven_multilingual_v2",
                output_format="mp3_44100_128"
            )
            
            # 스트림을 바이트로 변환
            if hasattr(audio_stream, 'read'):
                audio_bytes = audio_stream.read()
            elif isinstance(audio_stream, (bytes, bytearray)):
                audio_bytes = audio_stream
            else:
                audio_bytes = b''.join(chunk for chunk in audio_stream)
            
            # 파일 저장
            output_file = f"essay_audio_{datetime.now().strftime('%Y%m%d_%H%M%S')}.mp3"
            with open(output_file, "wb") as f:
                f.write(audio_bytes)
            
            return output_file
            
        except Exception as e:
            print(f"Error generating audio: {e}")
            raise